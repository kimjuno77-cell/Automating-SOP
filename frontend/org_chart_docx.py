"""
Org Chart Word & HTML Document Generator
Generates professional documents containing the organizational chart
and detailed responsibility breakdown for each role.
"""
import os
import io
import urllib.parse
import requests
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict
import base64


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _get_depth(roles, idx, visited=None):
    if visited is None:
        visited = set()
    if idx in visited:
        return 0
    visited.add(idx)
    parent = roles[idx].get('parent_idx')
    if parent is None or parent < 0 or parent >= len(roles):
        return 0
    return 1 + _get_depth(roles, parent, visited)


def generate_org_chart_dot(roles):
    """Generate a Graphviz DOT string for the org chart."""
    if not roles:
        return ""
        
    dot = 'digraph G {\n'
    dot += '    rankdir=TB;\n'
    dot += '    splines=ortho;\n'
    dot += '    nodesep=1.0;\n'
    dot += '    ranksep=1.0;\n'
    dot += '    dpi=150;\n'
    dot += '    node [shape=plaintext, fontname="Malgun Gothic", fontsize=10];\n'
    dot += '    edge [color="#334155", penwidth=1.5];\n'
    
    for i, role in enumerate(roles):
        name = role.get('role_name', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        desc = role.get('role_desc', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        resp_lines = role.get('responsibilities', '').split('\n')
        
        html_label = '<<TABLE BORDER="2" CELLBORDER="0" CELLSPACING="0" CELLPADDING="6" COLOR="#1e3a5f">\n'
        html_label += f'  <TR><TD BGCOLOR="#1e3a5f"><FONT COLOR="white"><B><U>{name}</U></B></FONT></TD></TR>\n'
        if desc.strip():
            html_label += f'  <TR><TD ALIGN="CENTER" BGCOLOR="#f0f4f8">{desc}</TD></TR>\n'
        
        if any(r.strip() for r in resp_lines):
            html_label += '  <TR><TD ALIGN="LEFT" BGCOLOR="white">'
            for r in resp_lines:
                if r.strip():
                    html_label += f'{r.strip()}<BR ALIGN="LEFT"/>'
            html_label += '</TD></TR>\n'
            
        html_label += '</TABLE>>'
        
        dot += f'    node_{i} [label={html_label}];\n'
    
    children_by_parent = defaultdict(list)
    for i, role in enumerate(roles):
        parent_idx = role.get('parent_idx')
        if parent_idx is not None and 0 <= parent_idx < len(roles) and parent_idx != i:
            dot += f'    node_{parent_idx} -> node_{i};\n'
            children_by_parent[parent_idx].append(i)
    
    for parent_idx, child_indices in children_by_parent.items():
        if len(child_indices) > 1:
            nodes_str = '; '.join(f'node_{c}' for c in child_indices)
            dot += f'    {{ rank=same; {nodes_str}; }}\n'
            
    dot += '}\n'
    return dot


def _fetch_chart_image(roles):
    """Fetch the org chart image from QuickChart using POST (no URL length limit)."""
    dot = generate_org_chart_dot(roles)
    if not dot:
        return None
    try:
        # Use POST to avoid URL length limits with complex org charts
        resp = requests.post(
            "https://quickchart.io/graphviz",
            json={"graph": dot, "format": "png"},
            timeout=30
        )
        if resp.status_code == 200:
            img_bytes = resp.content
            img = PILImage.open(io.BytesIO(img_bytes))
            w, h = img.size
            return img_bytes, w, h
    except Exception:
        pass
    return None


def generate_org_docx(roles, title="3.0 조직 및 업무분장"):
    """Generate a Word document for the Org Chart."""
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # Calculate usable width in inches
    usable_width_emu = section.page_width - section.left_margin - section.right_margin
    usable_width_inches = usable_width_emu / 914400  # EMU to inches
    usable_height_emu = section.page_height - section.top_margin - section.bottom_margin
    usable_height_inches = usable_height_emu / 914400
    
    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Generate and embed the org chart image — fit to page
    result = _fetch_chart_image(roles)
    if result:
        img_bytes, img_w_px, img_h_px = result
        image_stream = io.BytesIO(img_bytes)
        
        # Available space on page
        max_w = usable_width_inches
        max_h = usable_height_inches - 1.5  # room for title
        
        # Convert pixel dimensions to inches (at 150 DPI as set in DOT)
        img_w_in = img_w_px / 150.0
        img_h_in = img_h_px / 150.0
        
        # Calculate scale factors for both dimensions
        scale_w = max_w / img_w_in if img_w_in > 0 else 1
        scale_h = max_h / img_h_in if img_h_in > 0 else 1
        
        # Use the SMALLER scale to guarantee it fits BOTH dimensions
        scale = min(scale_w, scale_h, 1.0)  # also cap at 1.0 to not upscale
        
        fit_w = img_w_in * scale
        # Ensure fit_w never exceeds page width
        fit_w = min(fit_w, max_w)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_stream, width=Inches(fit_w))
    else:
        doc.add_paragraph("(조직도 이미지 생성 오류)")
    
    # Page break before the detail table
    doc.add_page_break()
    
    # Detailed Responsibility Table
    detail_title = doc.add_heading("업무분장 상세", level=2)
    
    # Build a table: 직책명 | 직책 설명 | 상위 직책 | 주요 업무
    table = doc.add_table(rows=1 + len(roles), cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ["직책명", "직책 설명", "상위 직책", "주요 업무"]
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "1e3a5f")
    
    # Data rows
    for i, role in enumerate(roles):
        row = table.rows[i + 1]
        
        # 직책명
        cell0 = row.cells[0]
        cell0.text = role.get('role_name', '')
        if cell0.paragraphs[0].runs:
            cell0.paragraphs[0].runs[0].font.size = Pt(9)
            cell0.paragraphs[0].runs[0].bold = True
        
        # 직책 설명
        cell1 = row.cells[1]
        cell1.text = role.get('role_desc', '')
        for run in cell1.paragraphs[0].runs:
            run.font.size = Pt(9)
        
        # 상위 직책
        cell2 = row.cells[2]
        pidx = role.get('parent_idx')
        if pidx is not None and 0 <= pidx < len(roles):
            cell2.text = roles[pidx].get('role_name', '')
        else:
            cell2.text = "(최상위)"
        for run in cell2.paragraphs[0].runs:
            run.font.size = Pt(9)
        
        # 주요 업무
        cell3 = row.cells[3]
        cell3.text = role.get('responsibilities', '')
        for run in cell3.paragraphs[0].runs:
            run.font.size = Pt(9)
        
        # Alternate row shading
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "f0f4f8")
    
    # Set column widths
    widths = [Cm(4), Cm(6), Cm(4), Cm(12)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_org_html(roles, title="3.0 조직 및 업무분장"):
    """Generate a professional HTML report for the Org Chart."""
    
    # Fetch chart image and convert to base64 for embedding
    img_b64 = ""
    result = _fetch_chart_image(roles)
    if result:
        img_bytes, _, _ = result
        img_b64 = base64.b64encode(img_bytes).decode('utf-8')
    
    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700&display=swap');
        
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        
        body {{
            font-family: 'Noto Sans KR', 'Malgun Gothic', sans-serif;
            background: #f8fafc;
            color: #1e293b;
            padding: 40px;
        }}
        
        .report-container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 12px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            overflow: hidden;
        }}
        
        .report-header {{
            background: linear-gradient(135deg, #1e3a5f 0%, #2563eb 100%);
            color: white;
            padding: 30px 40px;
        }}
        
        .report-header h1 {{
            font-size: 1.8rem;
            font-weight: 700;
            margin-bottom: 5px;
        }}
        
        .report-header p {{
            opacity: 0.85;
            font-size: 0.9rem;
        }}
        
        .section {{
            padding: 30px 40px;
        }}
        
        .section h2 {{
            font-size: 1.3rem;
            font-weight: 700;
            color: #1e3a5f;
            margin-bottom: 20px;
            padding-bottom: 8px;
            border-bottom: 2px solid #e2e8f0;
        }}
        
        .chart-container {{
            text-align: center;
            padding: 20px 0;
            overflow-x: auto;
        }}
        
        .chart-container img {{
            max-width: 100%;
            height: auto;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
            font-size: 0.9rem;
        }}
        
        thead th {{
            background: #1e3a5f;
            color: white;
            padding: 12px 16px;
            text-align: center;
            font-weight: 600;
        }}
        
        tbody td {{
            padding: 10px 16px;
            border-bottom: 1px solid #e2e8f0;
            vertical-align: top;
        }}
        
        tbody tr:nth-child(even) {{
            background: #f0f4f8;
        }}
        
        tbody tr:hover {{
            background: #e0f2fe;
        }}
        
        .role-name {{
            font-weight: 700;
            color: #1e3a5f;
        }}
        
        .responsibilities {{
            white-space: pre-line;
            font-size: 0.85rem;
            line-height: 1.6;
        }}
        
        .footer {{
            text-align: center;
            padding: 20px;
            color: #94a3b8;
            font-size: 0.8rem;
            border-top: 1px solid #e2e8f0;
        }}
        
        @media print {{
            body {{ padding: 0; background: white; }}
            .report-container {{ box-shadow: none; border-radius: 0; }}
            .chart-container {{ page-break-after: always; }}
        }}
    </style>
</head>
<body>
    <div class="report-container">
        <div class="report-header">
            <h1>{title}</h1>
            <p>EMKO SOP Automation System — 조직 및 업무분장 보고서</p>
        </div>
        
        <div class="section">
            <h2>📊 조직도</h2>
            <div class="chart-container">
"""
    
    if img_b64:
        html += f'                <img src="data:image/png;base64,{img_b64}" alt="조직도">\n'
    else:
        html += '                <p>(조직도 이미지를 생성할 수 없습니다.)</p>\n'
    
    html += """            </div>
        </div>
        
        <div class="section">
            <h2>📋 업무분장 상세</h2>
            <table>
                <thead>
                    <tr>
                        <th style="width:15%">직책명</th>
                        <th style="width:20%">직책 설명</th>
                        <th style="width:15%">상위 직책</th>
                        <th style="width:50%">주요 업무</th>
                    </tr>
                </thead>
                <tbody>
"""
    
    for i, role in enumerate(roles):
        name = role.get('role_name', '')
        desc = role.get('role_desc', '')
        resp = role.get('responsibilities', '').replace('\n', '<br>')
        pidx = role.get('parent_idx')
        if pidx is not None and 0 <= pidx < len(roles):
            parent_name = roles[pidx].get('role_name', '')
        else:
            parent_name = "(최상위)"
        
        html += f"""                    <tr>
                        <td class="role-name">{name}</td>
                        <td>{desc}</td>
                        <td>{parent_name}</td>
                        <td class="responsibilities">{resp}</td>
                    </tr>
"""
    
    html += """                </tbody>
            </table>
        </div>
        
        <div class="footer">
            © 2026 EMKO SOP Automation System | 조직 및 업무분장 보고서
        </div>
    </div>
</body>
</html>"""
    
    return html.encode('utf-8')
