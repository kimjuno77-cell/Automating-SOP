import streamlit as st
import io
import os
import copy
import base64
import requests
import urllib.parse
from PIL import Image as PILImage
from collections import defaultdict
from datetime import datetime
from org_chart import build_org_chart_url


# ============================================================
# SELF-CONTAINED document generators (no external module cache)
# ============================================================

def _get_dot(roles):
    """Generate Graphviz DOT string for the org chart."""
    if not roles:
        return ""
    dot = 'digraph G {\n'
    dot += '    rankdir=TB;\n'
    dot += '    splines=ortho;\n'
    dot += '    nodesep=1.0;\n'
    dot += '    ranksep=1.0;\n'
    dot += '    node [shape=plaintext, fontname="Malgun Gothic", fontsize=10];\n'
    dot += '    edge [color="#334155", penwidth=1.5];\n'
    for i, role in enumerate(roles):
        name = role.get('role_name', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        desc = role.get('role_desc', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        resp_lines = role.get('responsibilities', '').split('\n')
        html_label = '<<TABLE BORDER="2" CELLBORDER="0" CELLSPACING="0" CELLPADDING="6" COLOR="#1e3a5f">\n'
        html_label += f'  <TR><TD BGCOLOR="#1e3a5f"><FONT COLOR="white"><B><U>{name}</U></B></FONT></TD></TR>\n'
        if desc.strip():
            html_label += f'  <TR><TD ALIGN="CENTER" BGCOLOR="#f0f4f8">{desc}</TD></TR>\n'
        if any(r.strip() for r in resp_lines):
            html_label += '  <TR><TD ALIGN="LEFT" BGCOLOR="white">'
            for r in resp_lines:
                if r.strip():
                    html_label += f'{r.strip()}<BR ALIGN="LEFT"/>'
            html_label += '</TD></TR>\n'
        html_label += '</TABLE>>'
        dot += f'    node_{i} [label={html_label}];\n'
    children_by_parent = defaultdict(list)
    for i, role in enumerate(roles):
        parent_idx = role.get('parent_idx')
        if parent_idx is not None and 0 <= parent_idx < len(roles) and parent_idx != i:
            dot += f'    node_{parent_idx} -> node_{i};\n'
            children_by_parent[parent_idx].append(i)
    for parent_idx, child_indices in children_by_parent.items():
        if len(child_indices) > 1:
            nodes_str = '; '.join(f'node_{c}' for c in child_indices)
            dot += f'    {{ rank=same; {nodes_str}; }}\n'
    dot += '}\n'
    return dot


def _fetch_image_bytes(roles):
    """Fetch chart PNG from QuickChart POST API. Returns (img_bytes, width_px, height_px) or None."""
    dot = _get_dot(roles)
    if not dot:
        return None
    try:
        resp = requests.post(
            "https://quickchart.io/graphviz",
            json={"graph": dot, "format": "png"},
            timeout=30
        )
        if resp.status_code == 200:
            img_bytes = resp.content
            img = PILImage.open(io.BytesIO(img_bytes))
            w, h = img.size
            return img_bytes, w, h
    except Exception as e:
        st.error(f"이미지 fetch 오류: {e}")
    return None


def _make_docx(roles, title="3.0 조직 및 업무분장"):
    """Build Word docx bytes with the org chart image sized to fit the page."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _shade(cell, color):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)

    doc = Document()
    section = doc.sections[0]
    # Landscape orientation
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Usable area in inches (EMU / 914400)
    usable_w = (section.page_width - section.left_margin - section.right_margin) / 914400
    usable_h = (section.page_height - section.top_margin - section.bottom_margin) / 914400

    # Title
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Fetch image
    result = _fetch_image_bytes(roles)
    if result:
        img_bytes, img_w_px, img_h_px = result

        # Image natural size in inches at 72 DPI (Graphviz default)
        img_w_in = img_w_px / 72.0
        img_h_in = img_h_px / 72.0

        # Space available (leave 0.5 inch for title)
        avail_w = usable_w
        avail_h = usable_h - 0.5

        # Scale to fit within BOTH dimensions
        scale = min(avail_w / img_w_in, avail_h / img_h_in, 1.0)
        fit_w = img_w_in * scale

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(fit_w))
    else:
        doc.add_paragraph("(조직도 이미지 생성 오류)")

    doc.add_page_break()
    doc.add_heading("업무분장 상세", level=2)

    table = doc.add_table(rows=1 + len(roles), cols=4)
    table.style = 'Table Grid'
    for j, hdr in enumerate(["직책명", "직책 설명", "상위 직책", "주요 업무"]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade(cell, "1e3a5f")

    for i, role in enumerate(roles):
        row = table.rows[i + 1]
        cells = row.cells
        cells[0].text = role.get('role_name', '')
        cells[1].text = role.get('role_desc', '')
        pidx = role.get('parent_idx')
        cells[2].text = roles[pidx].get('role_name', '') if pidx is not None and 0 <= pidx < len(roles) else "(최상위)"
        cells[3].text = role.get('responsibilities', '')
        for c in cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
        if i % 2 == 0:
            for c in cells:
                _shade(c, "f0f4f8")

    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = [Cm(4), Cm(6), Cm(4), Cm(12)][j]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _make_html(roles, title="3.0 조직 및 업무분장"):
    """Build HTML report with embedded base64 chart image."""
    img_b64 = ""
    result = _fetch_image_bytes(roles)
    if result:
        img_bytes, _, _ = result
        img_b64 = base64.b64encode(img_bytes).decode('utf-8')

    rows_html = ""
    for i, role in enumerate(roles):
        name = role.get('role_name', '')
        desc = role.get('role_desc', '')
        resp = role.get('responsibilities', '').replace('\n', '<br>')
        pidx = role.get('parent_idx')
        parent_name = roles[pidx].get('role_name', '') if pidx is not None and 0 <= pidx < len(roles) else "(최상위)"
        rows_html += f"<tr><td style='font-weight:700;color:#1e3a5f'>{name}</td><td>{desc}</td><td>{parent_name}</td><td style='white-space:pre-line'>{resp}</td></tr>\n"

    img_tag = f'<img src="data:image/png;base64,{img_b64}" style="max-width:100%;height:auto" alt="조직도">' if img_b64 else '<p>(이미지 없음)</p>'

    html = f"""<!DOCTYPE html><html lang="ko"><head><meta charset="UTF-8">
<title>{title}</title>
<style>
  body{{font-family:'Malgun Gothic',sans-serif;background:#f8fafc;padding:30px;color:#1e293b}}
  .box{{max-width:1200px;margin:auto;background:white;border-radius:10px;box-shadow:0 4px 20px rgba(0,0,0,.08);overflow:hidden}}
  .hdr{{background:linear-gradient(135deg,#1e3a5f,#2563eb);color:white;padding:24px 36px}}
  .hdr h1{{font-size:1.6rem;margin:0 0 4px}}
  .sec{{padding:24px 36px}}
  .sec h2{{font-size:1.2rem;color:#1e3a5f;border-bottom:2px solid #e2e8f0;padding-bottom:6px;margin-bottom:16px}}
  .chart{{text-align:center;overflow-x:auto;padding:12px 0}}
  table{{width:100%;border-collapse:collapse;font-size:.9rem}}
  thead th{{background:#1e3a5f;color:white;padding:10px 14px;text-align:center}}
  tbody td{{padding:9px 14px;border-bottom:1px solid #e2e8f0;vertical-align:top}}
  tbody tr:nth-child(even){{background:#f0f4f8}}
  tbody tr:hover{{background:#e0f2fe}}
  .ft{{text-align:center;padding:16px;color:#94a3b8;font-size:.8rem;border-top:1px solid #e2e8f0}}
  @media print{{body{{padding:0;background:white}}.box{{box-shadow:none;border-radius:0}}.chart{{page-break-after:always}}}}
</style></head><body>
<div class="box">
  <div class="hdr"><h1>{title}</h1><p>EMKO SOP — 조직 및 업무분장 보고서</p></div>
  <div class="sec"><h2>📊 조직도</h2><div class="chart">{img_tag}</div></div>
  <div class="sec"><h2>📋 업무분장 상세</h2>
    <table><thead><tr><th style="width:15%">직책명</th><th style="width:20%">직책 설명</th><th style="width:15%">상위 직책</th><th style="width:50%">주요 업무</th></tr></thead>
    <tbody>{rows_html}</tbody></table>
  </div>
  <div class="ft">© 2026 EMKO SOP Automation System</div>
</div></body></html>"""
    return html.encode('utf-8')


# ==========================================
# Preset Organizational Templates
# ==========================================
ORG_PRESETS = {
    "직접 입력": None,
    "환경사업본부 (기본)": [
        {"role_name": "환경사업본부사장", "role_desc": "환경사업본부 전체 관리 및 감독", "responsibilities": "", "parent_idx": None},
        {"role_name": "기술팀장", "role_desc": "팀 및 PROJECT 업무관리감독", "responsibilities": "- 설계업무 수행계획서 검토 및 승인\n- DESIGN PLAN 검토 및 승인\n- 설계 M/H 관리", "parent_idx": 0},
        {"role_name": "사업(영업)팀장", "role_desc": "사업 수주 및 영업 관리", "responsibilities": "- 사업수행 계획서 검토 및 승인\n- Project Master Schedule 관리", "parent_idx": 0},
        {"role_name": "품질보증팀장", "role_desc": "품질 관리 및 감독", "responsibilities": "- ISO 관련 업무 절차 관리\n- 설계 품질관리 이행 및 설계요원의 품질교육", "parent_idx": 0},
        {"role_name": "전계장팀장", "role_desc": "전기/계장 설계 관리", "responsibilities": "- 전기/계장 설계 업무 총괄\n- 관련 도면 검토 및 승인", "parent_idx": 0},
        {"role_name": "프로 (기술팀)", "role_desc": "기술팀 프로 엔지니어", "responsibilities": "- 설계 업무지시 및 관리\n- 설계 계산서 및 도면 검증\n- M/R 작성 및 검증", "parent_idx": 1},
        {"role_name": "엔지니어 (기술팀)", "role_desc": "기술팀 엔지니어", "responsibilities": "- 설계업무 수행\n- SPEC 작성\n- M/R 작성\n- V/P 검토", "parent_idx": 1},
        {"role_name": "프로 (사업팀)", "role_desc": "사업팀 프로", "responsibilities": "- 사업 관리 및 조정\n- 고객 커뮤니케이션", "parent_idx": 2},
        {"role_name": "프로 (전계장팀)", "role_desc": "전계장팀 프로", "responsibilities": "- 전기/계장 설계 수행\n- 관련 도면 작성 및 검증", "parent_idx": 4}
    ],
    "플랜트 설계센터": [
        {"role_name": "EM 팀장", "role_desc": "플랜트 설계센터 업무의 전체 관리 및 감독", "responsibilities": "", "parent_idx": None},
        {"role_name": "장치설계팀장", "role_desc": "팀 및 PROJECT 업무관리감독", "responsibilities": "- 설계업무 수행계획서 검토 및 승인\n- DESIGN PLAN 검토 및 승인", "parent_idx": 0},
        {"role_name": "기계설계팀장", "role_desc": "(각 팀업무 GUIDE참조)", "responsibilities": "", "parent_idx": 0},
        {"role_name": "발전설계팀장", "role_desc": "(각 팀업무 GUIDE참조)", "responsibilities": "", "parent_idx": 0},
        {"role_name": "LEAD ENGINEER", "role_desc": "PROJECT 업무 수행 및 관리", "responsibilities": "- 설계 업무지시 및 관리 감독\n- 설계 계산서 및 도면 검증\n- M/R 작성 및 검증\n- TBE/CBE 작성 및 검증\n- V/P 검토 및 승인\n- SPEC 작성 및 검증", "parent_idx": 1},
        {"role_name": "SUB-LEAD ENGINEER", "role_desc": "LEAD ENGINEER의 지시에 따라 PROJECT업무 수행", "responsibilities": "- 설계업무\n- SPEC작성\n- M/R 작성\n- TBE/CBE 작성\n- V/P 검토\n- 설계 계산서 작성", "parent_idx": 1},
        {"role_name": "ENGINEER", "role_desc": "LEAD ENGINEER 및 SUB L/E의 지시에따라 PROJECT 업무 수행", "responsibilities": "- 설계업무\n- SPEC작성\n- M/R 작성\n- TBE 작성\n- V/P 검토\n- 설계 계산서 작성", "parent_idx": 1}
    ],
}

ORG_PRESET_NAMES = list(ORG_PRESETS.keys())


def _get_parent_options(roles, current_idx):
    """Build parent selection options: list of (display_label, index) pairs."""
    options = [("없음 (최상위)", -1)]
    for j, r in enumerate(roles):
        if j != current_idx:
            # Show hierarchy level indicator
            depth = _get_depth(roles, j)
            indent = "  " * depth
            options.append((f"{indent}[{j+1}] {r['role_name']}", j))
    return options


def _get_depth(roles, idx, visited=None):
    """Calculate hierarchy depth for display indentation."""
    if visited is None:
        visited = set()
    if idx in visited:
        return 0
    visited.add(idx)
    parent = roles[idx].get('parent_idx')
    if parent is None or parent < 0 or parent >= len(roles):
        return 0
    return 1 + _get_depth(roles, parent, visited)


def render_tab2():
    # --- Template Selector ---
    st.markdown("#### 📑 조직도 템플릿 선택")
    tc1, tc2 = st.columns([2, 1])
    with tc1:
        if 'org_preset_current' not in st.session_state:
            st.session_state.org_preset_current = ORG_PRESET_NAMES[0]
        
        selected_preset = st.selectbox(
            "미리 구성된 조직도 템플릿을 선택하세요",
            ORG_PRESET_NAMES,
            index=ORG_PRESET_NAMES.index(st.session_state.org_preset_current),
            key="org_preset_selector"
        )
        
        if selected_preset != st.session_state.org_preset_current:
            st.session_state.org_preset_current = selected_preset
            preset_data = ORG_PRESETS.get(selected_preset)
            if preset_data is not None:
                st.session_state.org_roles = copy.deepcopy(preset_data)
                st.session_state.org_form_key = st.session_state.get('org_form_key', 0) + 1
            st.rerun()

    st.markdown("---")
    
    left_col, right_col = st.columns([2, 2])
    fk = st.session_state.get('org_form_key', 0)
    
    with left_col:
        st.markdown("### 📝 조직 및 업무분장 입력")
        
        # Add various types of roles
        add_cols = st.columns(3)
        with add_cols[0]:
            if st.button("👤 최상위 직책 추가", key="add_role_top", use_container_width=True):
                st.session_state.org_roles.insert(0, {
                    "role_name": "새 직책", "role_desc": "", 
                    "responsibilities": "", "parent_idx": None
                })
                st.session_state.org_form_key = fk + 1
                st.rerun()
        with add_cols[1]:
            if st.button("👥 팀장급 추가", key="add_team_leader", use_container_width=True):
                # Find the top-level role (index 0) and add a child
                st.session_state.org_roles.append({
                    "role_name": "새 팀장", "role_desc": "팀업무 관리감독",
                    "responsibilities": "", "parent_idx": 0
                })
                st.session_state.org_form_key = fk + 1
                st.rerun()
        with add_cols[2]:
            if st.button("🧑‍💻 팀원 추가", key="add_member", use_container_width=True):
                st.session_state.org_roles.append({
                    "role_name": "새 팀원", "role_desc": "업무 수행",
                    "responsibilities": "", "parent_idx": None
                })
                st.session_state.org_form_key = fk + 1
                st.rerun()
            
        for i, role in enumerate(st.session_state.org_roles):
            depth = _get_depth(st.session_state.org_roles, i)
            level_icons = ["🔹", "🔸", "▫️"]
            icon = level_icons[min(depth, len(level_icons)-1)]
            indent_str = "&nbsp;" * (depth * 4)
            
            with st.expander(f"{icon} **{i+1}.** {role['role_name']}", expanded=False):
                ctrl = st.columns(4)
                with ctrl[0]:
                    if st.button("⬆️", key=f"rup_{i}_{fk}", help="위로 이동", disabled=(i == 0)):
                        st.session_state.org_roles[i], st.session_state.org_roles[i-1] = st.session_state.org_roles[i-1], st.session_state.org_roles[i]
                        st.rerun()
                with ctrl[1]:
                    if st.button("⬇️", key=f"rdown_{i}_{fk}", help="아래로 이동", disabled=(i >= len(st.session_state.org_roles) - 1)):
                        st.session_state.org_roles[i], st.session_state.org_roles[i+1] = st.session_state.org_roles[i+1], st.session_state.org_roles[i]
                        st.rerun()
                with ctrl[2]:
                    if st.button("➕ 하위 추가", key=f"rins_{i}_{fk}", help="이 직책 아래에 새 하위 직책 추가"):
                        st.session_state.org_roles.insert(i+1, {
                            "role_name": "새 직책", "role_desc": "", 
                            "responsibilities": "", "parent_idx": i
                        })
                        st.session_state.org_form_key = fk + 1
                        st.rerun()
                with ctrl[3]:
                    if st.button("🗑️ 삭제", key=f"rdel_{i}_{fk}", help="이 직책 삭제"):
                        st.session_state.org_roles.pop(i)
                        # Fix parent references
                        for r in st.session_state.org_roles:
                            if r.get('parent_idx') is not None:
                                if r['parent_idx'] == i:
                                    r['parent_idx'] = None
                                elif r['parent_idx'] > i:
                                    r['parent_idx'] -= 1
                        st.session_state.org_form_key = fk + 1
                        st.rerun()
                    
                role['role_name'] = st.text_input("직책명", role['role_name'], key=f"rn_{i}_{fk}")
                role['role_desc'] = st.text_input("직책 설명", role['role_desc'], key=f"rd_{i}_{fk}")
                
                # Parent selector with dropdown (much better UX than number input)
                parent_options = _get_parent_options(st.session_state.org_roles, i)
                parent_labels = [opt[0] for opt in parent_options]
                parent_values = [opt[1] for opt in parent_options]
                
                current_parent = role.get('parent_idx')
                if current_parent is None:
                    current_parent = -1
                try:
                    current_select_idx = parent_values.index(current_parent)
                except ValueError:
                    current_select_idx = 0
                
                selected_parent_label = st.selectbox(
                    "⬆️ 상위 직책 (같은 상위 = 평행관계)",
                    parent_labels,
                    index=current_select_idx,
                    key=f"rp_{i}_{fk}"
                )
                
                selected_parent_val = parent_values[parent_labels.index(selected_parent_label)]
                role['parent_idx'] = selected_parent_val if selected_parent_val >= 0 else None
                
                role['responsibilities'] = st.text_area(
                    "주요 업무 (줄바꿈으로 구분)", 
                    role['responsibilities'], height=150, key=f"rr_{i}_{fk}"
                )

    with right_col:
        st.markdown("### 👁️ 조직도 미리보기")
        st.caption("💡 같은 상위 직책을 가진 항목들은 자동으로 **평행(가로) 배치**됩니다.")
        if st.session_state.org_roles:
            try:
                # Use the inline POST method instead of the GET URL to avoid length limits
                result = _fetch_image_bytes(st.session_state.org_roles)
                if result:
                    img_bytes, w, h = result
                    st.image(img_bytes, use_container_width=True, caption="현재 조직도 구조 (POST)")
                else:
                    st.warning("순서도 생성 중 오류가 발생했습니다.")
            except Exception as e:
                st.error(f"미리보기 오류: {e}")
        else:
            st.info("직책을 추가하면 조직도가 자동으로 생성됩니다.")
        
        st.markdown("---")
        st.markdown("##### 📊 현재 조직 구조")
        for i, role in enumerate(st.session_state.org_roles):
            depth = _get_depth(st.session_state.org_roles, i)
            indent = "&nbsp;&nbsp;&nbsp;&nbsp;" * depth
            parent_name = ""
            pidx = role.get('parent_idx')
            if pidx is not None and 0 <= pidx < len(st.session_state.org_roles):
                parent_name = f" ← {st.session_state.org_roles[pidx]['role_name']}"
            st.markdown(f"{indent}**{i+1}.** {role['role_name']}{parent_name}", unsafe_allow_html=True)

    # =========================================
    # BOTTOM: Word & HTML Document Export
    # =========================================
    st.markdown("---")
    st.markdown("### 📄 조직도 & 업무분장 문서 출력")
    
    if 'org_file_generated' not in st.session_state:
        st.session_state.org_file_generated = False
    
    default_download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    
    if st.button("📥 조직도 & 업무분장 보고서 생성 (Word + HTML)", key="gen_org_docx", use_container_width=True, type="primary"):
        with st.spinner("조직도 보고서를 생성하고 있습니다... (순서도 이미지 렌더링 중)"):
            try:
                from datetime import datetime
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Determine writable directory
                try:
                    os.makedirs(default_download_dir, exist_ok=True)
                    output_dir = default_download_dir
                except Exception:
                    output_dir = "outputs"
                    os.makedirs(output_dir, exist_ok=True)
                
                # Generate Word
                docx_bytes = _make_docx(st.session_state.org_roles)
                docx_filename = f"조직_및_업무분장_{ts}.docx"
                docx_save_path = os.path.join(output_dir, docx_filename)
                
                # Generate HTML
                html_bytes = _make_html(st.session_state.org_roles)
                html_filename = f"조직_및_업무분장_{ts}.html"
                html_save_path = os.path.join(output_dir, html_filename)
                
                # Write with retry on PermissionError
                for path, data in [(docx_save_path, docx_bytes), (html_save_path, html_bytes)]:
                    try:
                        with open(path, "wb") as f:
                            f.write(data)
                    except PermissionError:
                        # File is locked (open in another app), try with suffix
                        base, ext = os.path.splitext(path)
                        for n in range(1, 10):
                            alt_path = f"{base}_{n}{ext}"
                            try:
                                with open(alt_path, "wb") as f:
                                    f.write(data)
                                if ext == '.docx':
                                    docx_save_path = alt_path
                                    docx_filename = os.path.basename(alt_path)
                                else:
                                    html_save_path = alt_path
                                    html_filename = os.path.basename(alt_path)
                                break
                            except PermissionError:
                                continue
                
                st.session_state.org_file_generated = True
                st.session_state.org_save_dir = output_dir
                st.session_state.org_generated_content = docx_bytes
                st.session_state.org_generated_filename = docx_filename
                st.session_state.org_html_content = html_bytes
                st.session_state.org_html_filename = html_filename
                
                st.toast("🎉 조직도 보고서 (Word + HTML) 생성 완료!", icon="✅")
                st.rerun()
            except Exception as e:
                st.error(f"문서 생성 중 오류: {e}")
    
    if st.session_state.get('org_file_generated', False):
        save_dir = st.session_state.get('org_save_dir', default_download_dir)
        st.success(
            f"🎉 **조직도 보고서가 생성 및 저장되었습니다!**\n\n"
            f"*   **💾 저장 폴더:** `{save_dir}`\n"
            f"*   **📄 Word 파일:** `{st.session_state.org_generated_filename}`\n"
            f"*   **🌐 HTML 파일:** `{st.session_state.org_html_filename}`"
        )
        
        btn1, btn2 = st.columns(2)
        with btn1:
            if st.button("📂 다운로드 폴더 열기", key="open_org_folder", use_container_width=True):
                try:
                    os.startfile(save_dir)
                except Exception:
                    pass
        
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button(
                label="📥 Word 재다운로드",
                data=st.session_state.org_generated_content,
                file_name=st.session_state.org_generated_filename,
                use_container_width=True,
                key="download_org_docx"
            )
        with dl2:
            st.download_button(
                label="📥 HTML 재다운로드",
                data=st.session_state.org_html_content,
                file_name=st.session_state.org_html_filename,
                use_container_width=True,
                key="download_org_html"
            )

