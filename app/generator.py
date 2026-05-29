import os
import requests
import urllib.parse
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .schemas import UnitMap, UnitMapStep

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_margins(cell, **kwargs):
    """
    Set cell padding (margins) in dxa (1/20th of a pt)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ('top', 'left', 'bottom', 'right'):
        val = kwargs.get(margin)
        if val is not None:
            node = OxmlElement('w:{}'.format(margin))
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def generate_full_flowchart_dot(unit_map: UnitMap) -> str:
    """
    Generates a beautiful, high-fidelity Graphviz DOT graph for the entire Unit Map.
    Features a two-column swimlane:
    - Left column: "고객/관련문서" (Customer/Related Document)
    - Right column: "업무팀" (Task Team)
    """
    steps = unit_map.steps
    if not steps:
        return ""
        
    dot = 'digraph G {\n'
    dot += '    rankdir=TB;\n'
    dot += '    splines=ortho;\n'
    dot += '    concentrate=true;\n'
    dot += '    nodesep=0.4;\n'
    dot += '    ranksep=0.3;\n'
    dot += '    bgcolor="transparent";\n'
    dot += '    pad=0.2;\n'
    
    # Global node and edge styles matching the premium look
    dot += '    node [fontname="Malgun Gothic", fontsize=10];\n'
    dot += '    edge [fontname="Malgun Gothic", fontsize=9];\n'
    
    # Define step nodes and related document nodes
    prev_step_id = None
    prev_global_doc_node = None
    
    for i, step in enumerate(steps):
        step_node = f"step_{i}"
        activity_clean = step.activity.replace('\n', '\\n').replace('"', '\\"')
        dept_clean = step.dept.replace('"', '\\"')
        
        # Define the right step box (solid, light blue filled, bold blue border)
        dot += f'    {step_node} [label="{activity_clean}\\n({dept_clean})", shape=box, style="filled,rounded", fillcolor="#e0f2fe", color="#0369a1", fontcolor="#0c4a6e", penwidth=2.5, width=2.2];\n'
        
        # If there are related documents for this step
        if step.related_docs:
            doc_nodes = []
            info_docs = [d for d in step.related_docs if d.classification == "정보제공"]
            review_docs = [d for d in step.related_docs if d.classification == "검토요청"]
            
            # 1. 정보제공 (Information Providing) - Left to Right
            for j, doc in enumerate(info_docs):
                doc_node = f"doc_info_{i}_{j}"
                doc_label = doc.label.replace('\n', '\\n').replace('"', '\\"')
                doc_dept = doc.department.replace('"', '\\"')
                
                # Define document node (dashed, white filled, gray border)
                dot += f'    {doc_node} [label="{doc_label}\\n({doc_dept})", shape=box, style="dashed,filled", fillcolor="#f8fafc", color="#64748b", fontcolor="#334155", penwidth=1.5, width=2.0];\n'
                doc_nodes.append(doc_node)
                
            # 2. 검토요청 (Review Request) - Right to Left
            for j, doc in enumerate(review_docs):
                doc_node = f"doc_rev_{i}_{j}"
                doc_label = doc.label.replace('\n', '\\n').replace('"', '\\"')
                doc_dept = doc.department.replace('"', '\\"')
                
                # Define document node (dashed, white filled, gray border)
                dot += f'    {doc_node} [label="{doc_label}\\n({doc_dept})", shape=box, style="dashed,filled", fillcolor="#f8fafc", color="#64748b", fontcolor="#334155", penwidth=1.5, width=2.0];\n'
                doc_nodes.append(doc_node)
                
            # Align documents vertically and force them on the left rank
            if doc_nodes:
                # Add invisible vertical connections to stack ALL of them on the left across steps
                for doc_node in doc_nodes:
                    if prev_global_doc_node is not None:
                        dot += f'    {prev_global_doc_node} -> {doc_node} [style=invis];\n'
                    prev_global_doc_node = doc_node
                
                # Align the middle document node with the step
                mid_idx = len(doc_nodes) // 2
                dot += f'    {{ rank=same; {doc_nodes[mid_idx]}; {step_node}; }}\n'
                
                # Draw connections
                for doc_node in doc_nodes:
                    if "info" in doc_node:
                        # 정보제공: Left to Right (solid blue arrow)
                        dot += f'    {doc_node} -> {step_node} [color="#0284c7", penwidth=1.5];\n'
                    else:
                        # 검토요청: Right to Left (dashed gray arrow)
                        dot += f'    {step_node} -> {doc_node} [style=dashed, color="#64748b", penwidth=1.5];\n'
                        
        # Connect steps sequentially on the right
        if prev_step_id is not None:
            dot += f'    {prev_step_id} -> {step_node} [color="#334155", penwidth=1.5, arrowsize=0.8];\n'
            
        prev_step_id = step_node
        
    dot += '}\n'
    return dot

def generate_full_flowchart_online(unit_map: UnitMap, output_path: str):
    dot_code = generate_full_flowchart_dot(unit_map)
    encoded_dot = urllib.parse.quote(dot_code.strip())
    url = f"https://quickchart.io/graphviz?format=png&graph={encoded_dot}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Diagram generation failed: {e}")
    return None

def generate_step_flowchart_dot(step: UnitMapStep, step_idx: int, total_steps: int) -> str:
    """
    Generates a high-fidelity Graphviz DOT graph for a single step.
    Aligns perfectly with table rows by keeping step_node in the center right,
    and documents stacked vertically on the left.
    Uses transparent background for seamless blending.
    """
    dot = 'digraph G {\n'
    dot += '    rankdir=TB;\n'
    dot += '    splines=true;\n'
    dot += '    bgcolor="transparent";\n'
    dot += '    pad=0.1;\n'
    dot += '    nodesep=0.25;\n'
    dot += '    ranksep=0.15;\n'
    dot += '    node [fontname="Malgun Gothic", fontsize=10];\n'
    dot += '    edge [fontname="Malgun Gothic", fontsize=9];\n'
    
    # Step process box (filled light blue, solid bold blue border)
    activity_clean = step.activity.replace('\n', '\\n').replace('"', '\\"')
    dept_clean = step.dept.replace('"', '\\"')
    dot += f'    step_node [label="{activity_clean}\\n({dept_clean})", shape=box, style="filled,rounded", fillcolor="#e0f2fe", color="#0369a1", fontcolor="#0c4a6e", penwidth=2.5, width=2.5];\n'
    
    # Vertical continuation anchors
    dot += '    top_anchor [shape=point, style=invis, width=0, height=0];\n'
    dot += '    bottom_anchor [shape=point, style=invis, width=0, height=0];\n'
    
    # In Step 0, render "고객/관련문서" and "업무팀" headers at the top of the cell
    if step_idx == 0:
        dot += '    hdr_left [label="고객/관련문서", shape=plaintext, fontname="Malgun Gothic Bold", fontsize=11, fontcolor="#475569"];\n'
        dot += '    hdr_right [label="업무팀", shape=plaintext, fontname="Malgun Gothic Bold", fontsize=11, fontcolor="#475569"];\n'
        dot += '    { rank=same; hdr_left; hdr_right; }\n'
        dot += '    hdr_left -> hdr_right [style=invis];\n'
        dot += '    hdr_right -> top_anchor [style=invis];\n'
        
    # Vertical connection lines
    if step_idx > 0:
        dot += '    top_anchor -> step_node [color="#334155", penwidth=1.5, arrowsize=0.8];\n'
    else:
        dot += '    top_anchor -> step_node [style=invis];\n'
        
    if step_idx < total_steps - 1:
        dot += '    step_node -> bottom_anchor [color="#334155", penwidth=1.5, arrowhead=none];\n'
    else:
        dot += '    step_node -> bottom_anchor [style=invis];\n'
        
    # Related documents on the left
    if step.related_docs:
        doc_nodes = []
        for j, doc in enumerate(step.related_docs):
            doc_node = f"doc_{j}"
            doc_label = doc.label.replace('\n', '\\n').replace('"', '\\"')
            doc_dept = doc.department.replace('"', '\\"')
            dot += f'    {doc_node} [label="{doc_label}\\n({doc_dept})", shape=box, style="dashed,filled", fillcolor="#f8fafc", color="#64748b", fontcolor="#334155", penwidth=1.5, width=2.2];\n'
            doc_nodes.append(doc_node)
            
        # Stack multiple documents vertically
        if len(doc_nodes) > 1:
            for k in range(len(doc_nodes) - 1):
                dot += f'    {doc_nodes[k]} -> {doc_nodes[k+1]} [style=invis];\n'
                
        # Center the stack with step node
        mid_idx = len(doc_nodes) // 2
        dot += f'    {{ rank=same; {doc_nodes[mid_idx]}; step_node; }}\n'
        
        # Connections using dir=back for review request (keeps it left-to-right)
        for doc_node, doc in zip(doc_nodes, step.related_docs):
            if doc.classification == "정보제공":
                dot += f'    {doc_node} -> step_node [color="#0284c7", penwidth=1.5];\n'
            else:
                dot += f'    {doc_node} -> step_node [style=dashed, color="#64748b", penwidth=1.5, dir=back];\n'
    else:
        # Dummy left anchor to preserve horizontal column layout
        dot += '    left_anchor [shape=point, style=invis, width=0, height=0];\n'
        dot += '    { rank=same; left_anchor; step_node; }\n'
        
    dot += '}\n'
    return dot

def generate_step_flowchart_online(step: UnitMapStep, step_idx: int, total_steps: int, output_path: str):
    dot_code = generate_step_flowchart_dot(step, step_idx, total_steps)
    encoded_dot = urllib.parse.quote(dot_code.strip())
    url = f"https://quickchart.io/graphviz?format=png&graph={encoded_dot}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Step diagram generation failed: {e}")
    return None

def generate_unit_map_docx(unit_map: UnitMap, output_path: str, temp_dir: str):

    doc = Document()
    
    # Set page to Landscape orientation with 0.75" margins (maximizing space)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Title Section (3 columns: Left = Ref No, Center = Title, Right = Hidden spacer for centering)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    # Set explicit widths for landscape layout (Total width = 9.5 Inches)
    widths = [Inches(1.5), Inches(6.5), Inches(1.5)]
    for i, col in enumerate(header_table.columns):
        col.width = widths[i]
        for cell in col.cells:
            cell.width = widths[i]
            
    hdr_cells = header_table.rows[0].cells
    
    # Left: Ref No
    ref_p = hdr_cells[0].paragraphs[0]
    ref_run = ref_p.add_run(f" {unit_map.ref_no} ")
    ref_run.bold = True
    ref_run.font.size = Pt(11)
    ref_run.font.name = "Malgun Gothic"
    ref_p.alignment = 1 # Center
    set_cell_border(hdr_cells[0], top={'val': 'single', 'sz': 12}, bottom={'val': 'single', 'sz': 12}, left={'val': 'single', 'sz': 12}, right={'val': 'single', 'sz': 12})
    set_cell_margins(hdr_cells[0], top=120, bottom=120, left=150, right=150)
    hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Center: Title
    title_p = hdr_cells[1].paragraphs[0]
    title_run = title_p.add_run(f" 단위업무 Map - {unit_map.title} ")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Malgun Gothic"
    title_p.alignment = 1 # Center
    set_cell_border(hdr_cells[1], top={'val': 'double', 'sz': 6}, bottom={'val': 'double', 'sz': 6}, left={'val': 'double', 'sz': 6}, right={'val': 'double', 'sz': 6})
    set_cell_margins(hdr_cells[1], top=120, bottom=120, left=150, right=150)
    hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Right: Empty (Hidden borders to maintain alignment)
    hdr_cells[2].text = ""
    set_cell_border(hdr_cells[2], top={'val': 'none'}, bottom={'val': 'none'}, left={'val': 'none'}, right={'val': 'none'})

    doc.add_paragraph() # Spacer

    # Main Table (7 visual columns: 2 for Process + 5 for data)
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    
    col_widths = [Inches(1.5), Inches(1.5), Inches(2.0), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.3)]
    
    # Set widths of table columns
    for i, col in enumerate(table.columns):
        col.width = col_widths[i]
        
    # Header Row 1: "업무 Process" merged across cols 0-1, other headers with rowspan=2
    hdr_row1 = table.rows[0].cells
    # Merge cols 0+1 for "업무 Process"
    process_cell = hdr_row1[0].merge(hdr_row1[1])
    p = process_cell.paragraphs[0]
    run = p.add_run("업무 Process")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = "Malgun Gothic"
    p.alignment = 1
    set_cell_margins(process_cell, top=100, bottom=100, left=100, right=100)
    process_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Other headers in row 1 (will be vertically merged with row 2)
    headers_right = ["업무처리요점", "시기/주기", "Input", "Output", "관련문서/기준"]
    for i, h in enumerate(headers_right):
        col_idx = i + 2
        cell = hdr_row1[col_idx]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Malgun Gothic"
        p.alignment = 1
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Header Row 2: sub-headers "고객/관련문서" | "업무팀"
    hdr_row2 = table.rows[1].cells
    sub_headers = ["고객/관련문서", "업무팀"]
    for i, sh in enumerate(sub_headers):
        cell = hdr_row2[i]
        p = cell.paragraphs[0]
        run = p.add_run(sh)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = "Malgun Gothic"
        p.alignment = 1
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Vertically merge cols 2-6 (header row 1 + header row 2)
    for col_idx in range(2, 7):
        table.cell(0, col_idx).merge(table.cell(1, col_idx))

    # Generate FULL flowchart image (all steps combined)
    full_flowchart_path = os.path.join(temp_dir, "full_flowchart.png")
    generate_full_flowchart_online(unit_map, full_flowchart_path)
    
    total_steps = len(unit_map.steps)
    for step_idx, step in enumerate(unit_map.steps):
        row_cells = table.add_row().cells
        
        # Explicit cell widths for columns in row
        for i, cell in enumerate(row_cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # First step row: merge cols 0+1 and insert full flowchart with rowspan
        if step_idx == 0:
            # Merge cols 0 and 1 for this row
            merged_cell = row_cells[0].merge(row_cells[1])
            p = merged_cell.paragraphs[0]
            p.alignment = 1
            if os.path.exists(full_flowchart_path):
                try:
                    p.add_run().add_picture(full_flowchart_path, width=Inches(2.8))
                except Exception as pic_err:
                    print(f"Failed to add full flowchart picture: {pic_err}")
            # Make borders transparent (no top/bottom internal lines)
            set_cell_border(merged_cell, top={'val': 'none'}, bottom={'val': 'none'})
        else:
            # Subsequent rows: merge cols 0+1 and leave empty (flowchart spans from row 0)
            merged_cell = row_cells[0].merge(row_cells[1])
            # Clear any text
            merged_cell.text = ""
            # Make borders transparent
            set_cell_border(merged_cell, top={'val': 'none'}, bottom={'val': 'none'})
            
        # Last row: restore bottom border for the process cell
        if step_idx == total_steps - 1:
            set_cell_border(merged_cell, bottom={'val': 'single', 'sz': 6})
            
        # Columns 2 to 6: Text Fields
        fields = [step.points, step.timing, step.input_data, step.output_data, step.standards]
        for idx, text in enumerate(fields, start=2):
            row_cells[idx].text = text
            # Ensure proper font and size for readability
            p = row_cells[idx].paragraphs[0]
            if p.runs:
                run = p.runs[0]
                run.font.name = "Malgun Gothic"
                run.font.size = Pt(9.5)

    doc.save(output_path)
    return output_path

def generate_unit_map_html(unit_map: UnitMap, output_path: str):
    logo_base64 = ""
    # Find and load the EMKO logo
    for path in ("EMKO Logo.png", "../EMKO Logo.png", "d:/업무 절차서(SOP) 자동화/EMKO Logo.png"):
        if os.path.exists(path):
            try:
                with open(path, "rb") as lf:
                    logo_base64 = base64.b64encode(lf.read()).decode()
                break
            except Exception:
                pass
                
    logo_html = ""
    if logo_base64:
        logo_html = f'<img src="data:image/png;base64,{logo_base64}" alt="EMKO Logo" style="height: 35px; width: auto;" />'
        
    steps_html = ""
    total_steps = len(unit_map.steps)
    
    # Generate FULL flowchart image and base64-encode it
    full_flowchart_path = os.path.join(os.path.dirname(output_path), "full_flowchart_html_temp.png")
    generate_full_flowchart_online(unit_map, full_flowchart_path)
    full_base64 = ""
    if os.path.exists(full_flowchart_path):
        try:
            with open(full_flowchart_path, "rb") as f:
                full_base64 = base64.b64encode(f.read()).decode()
            os.remove(full_flowchart_path)
        except Exception:
            pass
    
    for idx, step in enumerate(unit_map.steps):
        if idx == 0:
            # First row: flowchart cell with rowspan spanning all step rows
            flowchart_td = f"""
            <td colspan="2" rowspan="{total_steps}" style="text-align: center; vertical-align: middle; background-color: #ffffff; padding: 8px; border-left: 1px solid #cbd5e1; border-right: 1px solid #cbd5e1; border-top: none; border-bottom: 1px solid #cbd5e1;">
                {"<img src='data:image/png;base64," + full_base64 + "' style='width: 100%; height: auto;' />" if full_base64 else ""}
            </td>
            """
            steps_html += f"""
            <tr>
                {flowchart_td}
                <td>{step.points.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td style="text-align: center;">{step.timing.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.input_data.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.output_data.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.standards.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
            </tr>
            """
        else:
            # Subsequent rows: no process column (already spanned)
            steps_html += f"""
            <tr>
                <td>{step.points.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td style="text-align: center;">{step.timing.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.input_data.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.output_data.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
                <td>{step.standards.replace(chr(92)+'n', '<br>').replace(chr(10), '<br>')}</td>
            </tr>
            """
        
    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <title>단위업무 Map - {unit_map.title}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
        
        body {{
            font-family: 'Malgun Gothic', 'Inter', sans-serif;
            background-color: #f1f5f9;
            margin: 0;
            padding: 40px 20px;
            color: #1e293b;
        }}
        .document-container {{
            max-width: 1200px;
            margin: 0 auto;
            background-color: #ffffff;
            border-radius: 12px;
            box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.1), 0 8px 10px -6px rgba(0, 0, 0, 0.1);
            padding: 40px;
            box-sizing: border-box;
            position: relative;
        }}
        .header-top {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 30px;
            width: 100%;
        }}
        .logo-box {{
            width: 150px;
            display: flex;
            justify-content: flex-start;
        }}
        .title-box {{
            border: 4px double #1e293b;
            padding: 10px 40px;
            font-weight: 800;
            font-size: 1.4rem;
            text-align: center;
            color: #0f172a;
            background-color: #ffffff;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            border-radius: 6px;
            margin: 0 auto;
        }}
        .ref-box {{
            border: 1.5px solid #475569;
            padding: 8px 16px;
            font-weight: 700;
            font-size: 0.95rem;
            color: #1e293b;
            background-color: #f8fafc;
            border-radius: 6px;
            box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
            text-align: center;
            width: 100px;
        }}
        .main-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            font-size: 0.9rem;
        }}
        .main-table th {{
            background-color: #f1f5f9;
            color: #1e293b;
            font-weight: 700;
            border: 1px solid #cbd5e1;
            padding: 12px;
            text-align: center;
        }}
        .main-table td {{
            border: 1px solid #cbd5e1;
            padding: 12px;
            vertical-align: middle;
            color: #334155;
            line-height: 1.5;
        }}
        .footer {{
            text-align: center;
            color: #64748b;
            font-size: 0.8rem;
            margin-top: 40px;
            border-top: 1px solid #e2e8f0;
            padding-top: 20px;
        }}
        @media print {{
            body {{
                background-color: #ffffff;
                padding: 0;
            }}
            .document-container {{
                box-shadow: none;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    <div class="document-container">
        <div class="header-top">
            <div class="logo-box">
                {logo_html}
            </div>
            <div class="title-box">
                단위업무 Map - {unit_map.title}
            </div>
            <div class="ref-box">
                {unit_map.ref_no}
            </div>
        </div>
        
        <table class="main-table">
            <thead>
                <tr>
                    <th colspan="2" style="width: 40%;">업무 Process</th>
                    <th rowspan="2" style="width: 18%;">업무처리요점</th>
                    <th rowspan="2" style="width: 9%;">시기/주기</th>
                    <th rowspan="2" style="width: 10%;">Input</th>
                    <th rowspan="2" style="width: 10%;">Output</th>
                    <th rowspan="2" style="width: 13%;">관련문서/기준</th>
                </tr>
                <tr>
                    <th style="width: 15%;">고객/관련문서</th>
                    <th style="width: 15%;">업무팀</th>
                </tr>
            </thead>
            <tbody>
                {steps_html}
            </tbody>
        </table>
        
        <div class="footer">
            © 2026 EMKO SOP Automation System | Premium HTML Report
        </div>
    </div>
</body>
</html>
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    return output_path
